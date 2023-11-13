"""
tools for llamaindex_tgbot
"""

import os
import re
from pprint import pprint
import json
from datetime import datetime
import docx
import docx2txt
from IPython.display import display, FileLink
import random
import string

from llama_index import (
    SimpleDirectoryReader,
    VectorStoreIndex,
    StorageContext,
    load_index_from_storage,
    # PromptHelper,
    # ResponseSynthesizer,
)
from llama_index.node_parser.extractors import (
    MetadataExtractor,
    TitleExtractor,
    QuestionsAnsweredExtractor,
    KeywordExtractor,
    SummaryExtractor,
    EntityExtractor,
)
from llama_index.text_splitter import TokenTextSplitter
from llama_index.node_parser import SimpleNodeParser
from llama_index import LLMPredictor, ServiceContext, KnowledgeGraphIndex
from llama_index.retrievers import VectorIndexRetriever
from llama_index.query_engine import RetrieverQueryEngine
from llama_index.prompts import PromptTemplate
from llama_index.graph_stores import SimpleGraphStore
from llama_index.llms import OpenAI
import deepl
import openai
from langchain.chat_models import ChatOpenAI


with open("private/api_codes.json", "r") as json_file:
    api_codes = json.load(json_file)
with open("docs/glossary.json", "r") as json_file:
    glossary = json.load(json_file)

apigpt = api_codes["GPT"]["job"]
os.environ["OPENAI_API_KEY"] = apigpt
openai.api_key = apigpt
apideepl = api_codes["DeepL"]["third_free"]


llm = ChatOpenAI(temperature=0, model_name="gpt-3.5-turbo", max_tokens=2048)  # type: ignore
llm_predictor = LLMPredictor(llm=llm)  #

service_context = ServiceContext.from_defaults(
    llm_predictor=llm_predictor, chunk_size_limit=4000
)


def translate_to_ru(text, glssry=glossary["toRU"], auth_key=apideepl):
    """_summary_

    Args:
        text (_type_): _description_
        glssry (_type_, optional): _description_. Defaults to glossary["toRU"].
        auth_key (_type_, optional): _description_. Defaults to apideepl.

    Returns:
        _type_: _description_
    """
    translator_to_ru = deepl.Translator(auth_key)
    if glssry:
        to_ru_glossary = translator_to_ru.create_glossary(
            "My glossary",
            source_lang="EN-US",
            target_lang="RU",
            entries=glssry,
        )
        answer = translator_to_ru.translate_text(
            text,
            source_lang="EN",
            target_lang="RU",
            glossary=to_ru_glossary,
        ).text  # type: ignore
    else:
        answer = translator_to_ru.translate_text(
            text,
            source_lang="EN",
            target_lang="RU",
        ).text  # type: ignore
    return answer


def translate_to_en(text, glssry=glossary["toEN"], auth_key=apideepl):
    """_summary_

    Args:
        text (_type_): _description_
        glssry (_type_, optional): _description_. Defaults to glossary["toEN"].
        auth_key (_type_, optional): _description_. Defaults to apideepl.

    Returns:
        _type_: _description_
    """
    translator_to_en = deepl.Translator(auth_key)
    if glssry:
        to_en_glossary = translator_to_en.create_glossary(
            "My glossary",
            source_lang="RU",
            target_lang="EN-US",
            entries=glssry,
        )
        answer = translator_to_en.translate_text(
            text, source_lang="RU", target_lang="EN-US", glossary=to_en_glossary
        ).text  # type: ignore
    else:
        answer = translator_to_en.translate_text(
            text, source_lang="RU", target_lang="EN-US"
        ).text  # type: ignore
    return answer


def build_engine(storage, debug=False, custom_prompt=False):
    """_summary_

    Args:
        storage (_type_): _description_
        custom_prompt (bool, optional): _description_. Defaults to False.

    Returns:
        _type_: _description_
    """
    storage_context = StorageContext.from_defaults(persist_dir=storage)  # grani_index
    index = load_index_from_storage(storage_context, service_context=service_context)
    index_version = datetime.fromtimestamp(
        os.path.getmtime(storage + "vector_store.json")
    ).strftime("%Y-%m-%d %H-%M-%S")
    if debug:
        print(f"{index_version=}")

    retriever = VectorIndexRetriever(index=index, similarity_top_k=2)  # type: ignore
    resp_mode = "tree_summarize"
    query_engine = RetrieverQueryEngine.from_args(
        retriever,
        response_mode=resp_mode,  # response_mode="compact" # 2 минуты работает # type: ignore
    )
    if custom_prompt:
        prompt_map = {
            "compact": "text_qa_template",
            "tree_summarize": "summary_template",
        }
        qa_prompt_tmpl = PromptTemplate(
            "Context information is below.\n"
            "---------------------\n"
            "{context_str}\n"
            "---------------------\n"
            "Given the context information and not prior knowledge, answer the query. "
            "If you use your prior knowledge - add text 'GENERAL_KNOWLEDGE:' in the beginning of your answer\n"
            "Query: {query_str}\n"
            "Answer: "
        )
        query_engine.update_prompts(
            {f"response_synthesizer:{prompt_map[resp_mode]}": qa_prompt_tmpl}
        )
    return query_engine


def is_question(text):
    """Define question-words in text and return text with "?" in the end

    Args:
        text (_type_): _description_

    Returns:
        _type_: _description_
    """
    question_words = [
        "как",
        "каких",
        "что",
        "почему",
        "где",
        "когда",
        "кто",
        "какой",
        "сколько",
        "можно",
        "ли",
    ]
    if text[-1] != "?":
        text_low = text.lower()
        cleaned_text = re.sub(r"[^\w\s]", "", text_low)
        words = cleaned_text.split()
        for qw in question_words:
            if qw in words:
                return text + "?"
    return text


def display_file_links(path):
    """_summary_

    Args:
        path (_type_): _description_
    """
    path_en = "private/" + path
    link_en = FileLink(path_en, result_html_prefix="English file: ")

    path_ru = path_en.replace("_EN", "").replace("EN_", "")
    link_ru = FileLink(path_ru, result_html_prefix="Russian file: ")
    display(link_en, link_ru)


def model0707(text, storage, translator_toEN, translator_toRU, service_context, w=120):
    """_summary_

    Args:
        text (_type_): _description_
        storage (_type_): _description_
        translator_toEN (_type_): _description_
        translator_toRU (_type_): _description_
        service_context (_type_): _description_
        w (int, optional): _description_. Defaults to 120.
    """
    # Load VectorStoreIndex
    storage_context = StorageContext.from_defaults(persist_dir=storage)  # grani_index
    index = load_index_from_storage(storage_context, service_context=service_context)

    retriever = VectorIndexRetriever(index=index, similarity_top_k=2)
    query_engine = RetrieverQueryEngine.from_args(
        retriever, response_mode="tree_summarize"
    )  # response_mode="compact" # долго работает - 2 минуты

    print(text, end=" -> ")
    text = translator_toEN.translate_text(text, target_lang="EN-US").text
    print(text)

    # start_time = time.time()
    response = query_engine.query(text)
    # print(f"Запрос выполнен за {int(time.time() - start_time)} секунд:")
    pprint(response.response.strip(), width=w)
    answer = translator_toRU.translate_text(
        response.response.strip(), target_lang="RU"
    ).text

    pprint(answer, width=w)
    print(response.source_nodes[0].node.text.split("Detailed information: ")[-1])
    print(response.source_nodes[0].node.extra_info["filename"])
    display_file_links(response.source_nodes[0].node.extra_info["filename"])


def model1004(query_ru, storage, w=120):
    """Old model

    Args:
        query_ru (_type_): _description_
        storage (_type_): _description_
        translator_toEN (_type_): _description_
        translator_toRU (_type_): _description_
        w (int, optional): _description_. Defaults to 120.
    """
    # Load VectorStoreIndex
    query_engine = build_engine(storage)

    query_ru = is_question(query_ru)

    print(query_ru, end=" -> ")

    query_en = translate_to_en(query_ru)
    print(query_en)

    answer_en = query_engine.query(query_en)
    used_node = answer_en.source_nodes[0].node.extra_info["filename"]
    used_node_score = answer_en.source_nodes[0].score
    url_links = re.findall(
        r"https?://[^\s/$.?#].[^\s]*", answer_en.source_nodes[0].node.text
    )

    pprint(answer_en.response.strip(), width=w)  # type: ignore

    answer_ru = translate_to_ru(answer_en.response.strip())  # type: ignore

    pprint(answer_ru, width=w)
    # pprint(abb_replace_EN(answer, replacement_EN), width=w)
    print(url_links)
    print(f'{used_node.split("/")[-1]}. Confidence: {used_node_score:.3f}')
    display_file_links(used_node)


def model1004s(query_ru, storage, w=120):
    """query_en = f'answer as accurately as possible - "{query_en}"?'

    Args:
        query_ru (_type_): _description_
        storage (_type_): _description_
        translator_toEN (_type_): _description_
        translator_toRU (_type_): _description_
        w (int, optional): _description_. Defaults to 120.
    """
    # Load VectorStoreIndex
    query_engine = build_engine(storage)

    query_ru = is_question(query_ru)

    print(query_ru, end=" -> ")

    query_en = translate_to_en(query_ru)
    query_en = f'answer as accurately as possible - "{query_en}"?'
    print(query_en)

    answer_en = query_engine.query(query_en)
    used_node = answer_en.source_nodes[0].node.extra_info["filename"]
    url_links = re.findall(
        r"https?://[^\s/$.?#].[^\s]*", answer_en.source_nodes[0].node.text
    )

    pprint(answer_en.response.strip(), width=w)  # type: ignore

    answer_ru = translate_to_ru(answer_en.response.strip())  # type: ignore

    pprint(answer_ru, width=w)
    # pprint(abb_replace_EN(answer, replacement_EN), width=w)
    if url_links:
        print(url_links)
    print(used_node.split("/")[-1])
    display_file_links(used_node)


def model1004kg(query_ru, storage, service_context, w=120):
    """KG model

    Args:
        query_ru (_type_): _description_
        storage (_type_): _description_
        translator_toEN (_type_): _description_
        translator_toRU (_type_): _description_
        w (int, optional): _description_. Defaults to 120.
    """
    # Load VectorStoreIndex
    storage_context = StorageContext.from_defaults(persist_dir=storage)  # grani_index
    index = load_index_from_storage(storage_context, service_context=service_context)

    query_engine = index.as_query_engine(
        include_text=False, response_mode="tree_summarize"
    )

    query_ru = is_question(query_ru)

    print(query_ru, end=" -> ")

    query_en = translate_to_en(query_ru)
    print(query_en)

    answer_en = query_engine.query(query_en)
    # used_node = answer_en.source_nodes[0].node.extra_info["filename"]
    url_links = re.findall(
        r"https?://[^\s/$.?#].[^\s]*", answer_en.source_nodes[0].node.text
    )

    pprint(answer_en.response.strip(), width=w)  # type: ignore

    answer_ru = translate_to_ru(answer_en.response.strip())  # type: ignore

    pprint(answer_ru, width=w)
    print(url_links)
    return answer_en


def model1107(query_ru, storage, debug=False, w=120):
    """Base model. Added custom prompt for detecting source of answer.

    Args:
        query_ru (_type_): _description_
        storage (_type_): _description_
        translator_toEN (_type_): _description_
        translator_toRU (_type_): _description_
        w (int, optional): _description_. Defaults to 120.
    """
    # Load Model
    query_engine = build_engine(storage, custom_prompt=True)

    query_ru = is_question(query_ru)
    query_en = translate_to_en(query_ru)

    if debug:
        print(f"{query_ru} -> {query_en}")
    answer_en = query_engine.query(query_en)

    # For custom prompt
    answer_en_text = answer_en.response.strip()  # type: ignore
    additional_info = ""
    if answer_en_text.find("GENERAL_KNOWLEDGE: ") == -1:
        answer_en_node = answer_en.source_nodes[0]
        used_node = answer_en_node.node.extra_info["filename"]
        used_node_score = answer_en_node.score
        url_links = re.findall(r"https?://[^\s/$.?#].[^\s]*", answer_en_node.node.text)
        additional_info += f"\nПодробная информация: {url_links[0]}"
        additional_info += (
            f'\n{used_node.split("/")[-1]}. Confidence: {used_node_score:.3f}'
        )
        if debug:
            # print(url_links)
            # print(f'{used_node.split("/")[-1]}. Confidence: {used_node_score:.3f}')
            display_file_links(used_node)
    else:
        answer_en_text = answer_en_text.replace(
            "GENERAL_KNOWLEDGE: ", "По данным ChatGPT: "
        )

    answer_ru = translate_to_ru(answer_en_text)
    answer_ru += additional_info
    if debug:
        pprint(answer_en_text, width=w)
        pprint(answer_ru, width=w)

    #### Logs
    current_time = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    random_chars = "".join(random.choices(string.ascii_letters + string.digits, k=8))
    file_name = current_time + "_" + random_chars + ".txt"
    with open("logs/" + file_name, "w") as file:
        file.write(str(query_ru) + "\n\n")
        file.write(str(answer_ru) + "\n\n")
        file.write(str(query_en) + "\n")
        file.write(str(answer_en) + "\n")
    ####
    return answer_ru


def model1107meta(query_ru, storage, debug=False, w=120):
    """Added meta information for files

    Args:
        query_ru (_type_): _description_
        storage (_type_): _description_
        translator_toEN (_type_): _description_
        translator_toRU (_type_): _description_
        w (int, optional): _description_. Defaults to 120.
    """
    # Load Model
    query_engine = build_engine(storage, debug=debug, custom_prompt=True)

    query_ru = is_question(query_ru)
    query_en = translate_to_en(query_ru)

    if debug:
        print(f"{query_ru} -> {query_en}")
    answer_en = query_engine.query(query_en)

    # For custom prompt
    answer_en_text = answer_en.response.strip()  # type: ignore

    additional_info = ""
    if answer_en_text.find("GENERAL_KNOWLEDGE: ") == -1:
        answer_en_node = answer_en.source_nodes[0]
        # added in meta
        filename = answer_en.source_nodes[0].node.metadata[
            "file_name"
        ]  # replaced in meta
        filepath = storage.replace("storage/", "").replace("_meta", "")
        used_node = ""
        for root, dirs, files in os.walk(filepath):
            if filename in files:
                used_node = os.path.join(root, filename)
        used_node = used_node.replace("private/", "")
        ###############
        used_node_score = answer_en_node.score
        url_links = re.findall(r"https?://[^\s/$.?#].[^\s]*", answer_en_node.node.text)
        additional_info += f"\nПодробная информация: {url_links[0]}"
        additional_info += (
            f'\n{used_node.split("/")[-1]}. Confidence: {used_node_score:.3f}'
        )
        if debug:
            # print(url_links)
            # print(f'{used_node.split("/")[-1]}. Confidence: {used_node_score:.3f}')
            display_file_links(used_node)
    else:
        answer_en_text = answer_en_text.replace(
            "GENERAL_KNOWLEDGE: ", "По данным ChatGPT: "
        )

    answer_ru = translate_to_ru(answer_en_text)
    answer_ru += additional_info
    if debug:
        pprint(answer_en_text, width=w)
        pprint(answer_ru, width=w)

    #### Logs
    current_time = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    random_chars = "".join(random.choices(string.ascii_letters + string.digits, k=8))
    file_name = current_time + "_" + random_chars + ".txt"
    with open("logs/" + file_name, "w") as file:
        file.write(str(query_ru) + "\n\n")
        file.write(str(answer_ru) + "\n\n")
        file.write(str(query_en) + "\n")
        file.write(str(answer_en) + "\n")
    ####
    return answer_ru


def translate_word_files(folder_path):
    """todo probably can be deleted

    Args:
        folder_path (_type_): _description_
    """
    new_folder_path = folder_path + "_EN"
    os.makedirs(
        new_folder_path, exist_ok=True
    )  # Create the new folder if it doesn't exist

    for root, _, files in os.walk(folder_path):
        for file in files:
            if file.endswith(".docx"):
                file_path = os.path.join(root, file)
                new_file_path = os.path.join(new_folder_path, "EN_" + file)

                doc = docx.Document(file_path)
                new_doc = docx.Document()
                for para in doc.paragraphs:
                    print(para)
                    if para.text:
                        output = translate_to_en(para.text)
                    else:
                        print(file_path)
                        output = "\n"
                    new_doc.add_paragraph(output)
                new_doc.add_paragraph("\nDetailed information: ")

                new_doc.save(
                    new_file_path,
                )


def translate_word_files2(folder_path, new_folder_path):
    """
    37k symbols
    """
    for root, _, files in os.walk(folder_path):
        for file in files:
            if file.endswith(".docx"):
                file_path = os.path.join(root, file)
                subfolder = new_folder_path + "/" + "/".join(root.split("/")[2:])
                os.makedirs(subfolder, exist_ok=True)
                new_file_path = subfolder + "/EN_" + file
                print(new_file_path)

                new_doc = docx.Document()
                doc = docx2txt.process(file_path)
                paragraphs = doc.split("\n\n")
                for para in paragraphs:
                    if para:
                        output = translate_to_en(para)
                    else:
                        output = "\n"
                    new_doc.add_paragraph(output)
                new_doc.save(
                    new_file_path,
                )


def build_index_vs(folder):
    """VectorStoreIndex. Saved index in "storage/" + folder + "/"

    Args:
        folder (str): Path with translated to english files ("docs/2023-11-01_EN/")
    """
    file_metadata = lambda x: {"filename": x}
    documents = SimpleDirectoryReader(
        folder, file_metadata=file_metadata, recursive=True
    ).load_data()
    index = VectorStoreIndex.from_documents(  # GPTVectorStoreIndex (GPTTreeIndex)
        documents,
        service_context=service_context,
    )
    index.storage_context.persist("storage/" + folder + "/")


def build_index_kg(folder):
    """KnowledgeGraphIndex. Saved index in "storage/" + folder[:-1] + "_KG/"

    Args:
        folder (str): Path with translated to english files ("docs/2023-11-01_EN/")
    """
    max_triplets = 2
    file_metadata = lambda x: {"filename": x}
    documents = SimpleDirectoryReader(
        folder,
        file_metadata=file_metadata,
        required_exts=[".docx"],
        recursive=True,
    ).load_data()
    graph_store = SimpleGraphStore()
    storage_context = StorageContext.from_defaults(graph_store=graph_store)
    index = KnowledgeGraphIndex.from_documents(
        documents,
        max_triplets_per_chunk=max_triplets,
        storage_context=storage_context,
        service_context=service_context,
        # include_embeddings=True,
    )
    index.storage_context.persist("storage/" + folder[:-1] + "_KG/")


def build_index_vs_meta(folder):
    """VectorStoreIndex with metainformation. Saved index in "storage/" + folder[:-1] + "_meta/"

    Args:
        folder (str): Path with translated to english files ("docs/2023-11-01_EN/")
    """

    text_splitter = TokenTextSplitter(separator=" ", chunk_size=512, chunk_overlap=128)
    metadata_extractor = MetadataExtractor(
        extractors=[
            TitleExtractor(nodes=5, llm=llm),
            KeywordExtractor(keywords=10, llm=llm),  # задаем количество ключевых слов
            # QuestionsAnsweredExtractor(questions=1), # ошибка cannot
            # SummaryExtractor(summaries=["prev", "self"], llm=llm), # ошибка cannot
            # EntityExtractor(prediction_threshold=0.5),
        ],
    )
    node_parser = SimpleNodeParser.from_defaults(
        text_splitter=text_splitter,
        metadata_extractor=metadata_extractor,
    )
    documents = SimpleDirectoryReader(
        folder,
        recursive=True,
    ).load_data()
    nodes_with_meta = node_parser.get_nodes_from_documents(documents)
    # print(nodes_with_meta[0].metadata) # 'file_name', 'document_title', 'excerpt_keywords'
    new_index = VectorStoreIndex(nodes_with_meta)
    new_index.storage_context.persist("storage/" + folder[:-1] + "_meta/")
